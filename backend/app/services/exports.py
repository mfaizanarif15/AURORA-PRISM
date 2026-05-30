from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from loguru import logger
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import get_settings
from app.models import ClipCandidate, Episode, ExportPack
from app.services.assets import safe_filename
from app.services.observability import observation, safe_update
from app.services.transcripts import seconds_to_timestamp


async def create_export_pack(session: AsyncSession, episode_id: str) -> ExportPack:
    logger.info("Creating export pack episode_id={}", episode_id)
    with observation(
        "create_export_pack",
        as_type="span",
        input={"episode_id": episode_id},
        metadata={"operation": "export"},
    ) as span:
        episode = await session.get(Episode, episode_id)
        if episode is None:
            logger.warning("Export pack failed, episode not found episode_id={}", episode_id)
            raise ValueError("Episode not found")

        pack = ExportPack(episode_id=episode_id, status="running", manifest={})
        session.add(pack)
        await session.flush()
        logger.info("Export pack record created episode_id={} export_pack_id={}", episode_id, pack.id)
        try:
            clips = await _approved_clips(session, episode_id)
            root = get_settings().exports_dir / episode_id / "handoff"
            if root.exists():
                logger.debug("Removing previous handoff directory path={}", root)
                shutil.rmtree(root)
            root.mkdir(parents=True, exist_ok=True)
            manifest = _manifest(episode, clips)
            logger.info("Writing export assets episode_id={} clip_count={} root={}", episode_id, len(clips), root)
            _write_markdown(root / "handoff.md", episode, clips)
            _write_csv(root / "clips.csv", clips)
            _write_pdf(root / "handoff.pdf", episode, clips)
            _write_docx(root / "handoff.docx", episode, clips)
            _copy_rendered(root / "media", clips)
            zip_path = (
                get_settings().exports_dir / episode_id / f"{safe_filename(episode.title)}-handoff.zip"
            )
            if zip_path.exists():
                logger.debug("Removing previous export zip path={}", zip_path)
                zip_path.unlink()
            _zip_dir(root, zip_path)
            pack.status = "completed"
            pack.path = str(zip_path)
            pack.filename = zip_path.name
            pack.manifest = manifest
            logger.info("Export pack completed episode_id={} export_pack_id={} zip_path={}", episode_id, pack.id, zip_path)
        except Exception as exc:
            pack.status = "failed"
            pack.error = str(exc)
            logger.exception(
                "Export pack failed episode_id={} export_pack_id={} error={}",
                episode_id,
                pack.id,
                exc,
            )
        await session.commit()
        await session.refresh(pack)
        safe_update(
            span,
            output={
                "export_pack_id": pack.id,
                "status": pack.status,
                "filename": pack.filename,
                "clip_count": pack.manifest.get("clip_count") if pack.manifest else None,
            },
        )
        return pack


async def _approved_clips(session: AsyncSession, episode_id: str) -> list[ClipCandidate]:
    result = await session.execute(
        select(ClipCandidate)
        .where(ClipCandidate.episode_id == episode_id, ClipCandidate.status.in_(["approved", "exported"]))
        .options(
            selectinload(ClipCandidate.score),
            selectinload(ClipCandidate.metadata_items),
            selectinload(ClipCandidate.rendered_clips),
        )
        .order_by(ClipCandidate.rank)
    )
    clips = list(result.scalars())
    if not clips:
        logger.info("No approved clips found, using top candidates episode_id={}", episode_id)
        result = await session.execute(
            select(ClipCandidate)
            .where(ClipCandidate.episode_id == episode_id)
            .options(
                selectinload(ClipCandidate.score),
                selectinload(ClipCandidate.metadata_items),
                selectinload(ClipCandidate.rendered_clips),
            )
            .order_by(ClipCandidate.rank)
            .limit(5)
        )
        clips = list(result.scalars())
    logger.debug("Selected clips for export episode_id={} count={}", episode_id, len(clips))
    return clips


def _manifest(episode: Episode, clips: list[ClipCandidate]) -> dict:
    return {
        "episode_id": episode.id,
        "title": episode.title,
        "clip_count": len(clips),
        "clips": [
            {
                "id": clip.id,
                "clip_type": clip.clip_type,
                "start": seconds_to_timestamp(clip.start_seconds),
                "end": seconds_to_timestamp(clip.end_seconds),
                "score": clip.score.total_score if clip.score else None,
            }
            for clip in clips
        ],
    }


def _write_markdown(path: Path, episode: Episode, clips: list[ClipCandidate]) -> None:
    logger.debug("Writing markdown export path={} clip_count={}", path, len(clips))
    lines = [f"# {episode.title} Handoff", "", f"Guest: {episode.guest_name or 'TBD'}", ""]
    for clip in clips:
        lines.extend(
            [
                f"## Clip {clip.rank}: {clip.clip_type.title()}",
                f"- Time: {seconds_to_timestamp(clip.start_seconds)} - {seconds_to_timestamp(clip.end_seconds)}",
                f"- Score: {clip.score.total_score if clip.score else 'N/A'}",
                f"- Moment: {clip.moment_type}",
                f"- Reasoning: {clip.reasoning}",
                "",
                clip.excerpt,
                "",
            ]
        )
        for item in clip.metadata_items:
            lines.extend(
                [
                    f"### {item.platform}",
                    f"Title: {item.title}",
                    f"Hook: {item.hook}",
                    f"Caption: {item.caption}",
                    f"CTA: {item.soft_cta}",
                    "",
                ]
            )
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_csv(path: Path, clips: list[ClipCandidate]) -> None:
    logger.debug("Writing CSV export path={} clip_count={}", path, len(clips))
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["rank", "clip_type", "status", "start", "end", "score", "excerpt"])
        for clip in clips:
            writer.writerow(
                [
                    clip.rank,
                    clip.clip_type,
                    clip.status,
                    seconds_to_timestamp(clip.start_seconds),
                    seconds_to_timestamp(clip.end_seconds),
                    clip.score.total_score if clip.score else "",
                    clip.excerpt,
                ]
            )


def _write_pdf(path: Path, episode: Episode, clips: list[ClipCandidate]) -> None:
    logger.debug("Writing PDF export path={} clip_count={}", path, len(clips))
    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 48
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(48, y, f"{episode.title} Handoff")
    y -= 30
    pdf.setFont("Helvetica", 10)
    for clip in clips:
        if y < 120:
            pdf.showPage()
            y = height - 48
            pdf.setFont("Helvetica", 10)
        pdf.drawString(
            48,
            y,
            f"Clip {clip.rank} | {clip.clip_type} | {seconds_to_timestamp(clip.start_seconds)} - {seconds_to_timestamp(clip.end_seconds)}",
        )
        y -= 16
        for line in _wrap(clip.reasoning, 95)[:4]:
            pdf.drawString(58, y, line)
            y -= 13
        y -= 8
    pdf.save()


def _write_docx(path: Path, episode: Episode, clips: list[ClipCandidate]) -> None:
    logger.debug("Writing DOCX export path={} clip_count={}", path, len(clips))
    document = Document()
    document.add_heading(f"{episode.title} Handoff", 0)
    document.add_paragraph(f"Guest: {episode.guest_name or 'TBD'}")
    for clip in clips:
        document.add_heading(f"Clip {clip.rank}: {clip.clip_type.title()}", level=1)
        document.add_paragraph(
            f"{seconds_to_timestamp(clip.start_seconds)} - {seconds_to_timestamp(clip.end_seconds)}"
        )
        document.add_paragraph(f"Score: {clip.score.total_score if clip.score else 'N/A'}")
        document.add_paragraph(clip.reasoning)
        document.add_paragraph(clip.excerpt)
    document.save(str(path))


def _copy_rendered(directory: Path, clips: list[ClipCandidate]) -> None:
    directory.mkdir(parents=True, exist_ok=True)
    copied = 0
    for clip in clips:
        for rendered in clip.rendered_clips:
            if rendered.status == "completed" and rendered.path and Path(rendered.path).exists():
                shutil.copy2(rendered.path, directory / Path(rendered.path).name)
                copied += 1
    logger.debug("Copied rendered media directory={} copied_count={}", directory, copied)


def _zip_dir(source: Path, target: Path) -> None:
    file_count = 0
    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in source.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(source))
                file_count += 1
    logger.debug("Created export zip source={} target={} file_count={}", source, target, file_count)


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line: list[str] = []
    for word in words:
        if sum(len(item) for item in line) + len(line) + len(word) > width:
            lines.append(" ".join(line))
            line = [word]
        else:
            line.append(word)
    if line:
        lines.append(" ".join(line))
    return lines
